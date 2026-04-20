from __future__ import annotations

from io import BytesIO

from docx import Document
from openpyxl import Workbook

from ultimate_search.models import EvidenceItem


def evidence_to_xlsx(rows: list[dict]) -> bytes:
    buffer = BytesIO()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Evidence"
    if rows:
        headers = list(rows[0].keys())
        worksheet.append(headers)
        for row in rows:
            worksheet.append([row.get(header, "") for header in headers])
        worksheet.freeze_panes = "A2"
        for column_cells in worksheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in column_cells)
            worksheet.column_dimensions[column_cells[0].column_letter].width = min(max_length + 2, 55)
    workbook.save(buffer)
    return buffer.getvalue()


def brief_to_docx(brief_markdown: str, evidence: list[EvidenceItem], question: str) -> bytes:
    document = Document()
    document.add_heading("Ultimate Search Brief", level=1)
    document.add_paragraph(question)

    for line in brief_markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped.replace("### ", ""), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(stripped.replace("**", ""))

    if evidence:
        document.add_heading("Evidence Links", level=2)
        for idx, item in enumerate(evidence, start=1):
            document.add_paragraph(f"{idx}. {item.title} - {item.url}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
